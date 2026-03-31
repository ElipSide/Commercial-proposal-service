import base64
import json
import os
import re
import uuid
import aiohttp
from dadata import Dadata
import docx
from PyPDF2 import PdfReader, PdfWriter


class DocumentParser:
   
    def __init__(self, path, extension, chat_id):
        self.SECRET_TOKEN = "NDA2MmU3MTUtMjRjYi00YWI1LTk4Y2YtNzFjYWY0Zjk4NGY2OmU0ZjZiNjI0LWFhNTctNDI1Mi05MjkyLTM3YTllZjM5ZmEwMA=="
        self.oauth_token = "y0_AgAAAAByU-qRAATuwQAAAAEVvvgjAAA-h95ul_RP8an3hlDnnjxhSmRYJA"
        self.folder_id = "b1g73p115jr2vpi0cjgf"
        self.path = path
        self.extension = extension
        self.chat_id = chat_id

    async def main(self):
        # Данные полученные из документа, изображения
        data_company = {'Наименование организации': '', 'ИНН': '', 'БИК': '', 'Расчетный счет': '', 'Корреспондентский счет': '', 'Телефон': '', 'Электронная почта': '', 'Основание актов': '', 'Город': '', 'Адрес': '', 'Имя': '', 'Фамилия': '', 'Отчество': ''}
        if self.extension == 'docx': # получаем текст из docx файла
            text = self.extract_text_from_docx(self.path) 
        elif self.extension == 'txt': # получаем текст из txt файла
            with open(self.path, 'r', encoding='utf-8') as file:
                text_file = file.read()
            text = text_file.strip()
        elif self.extension == 'pdf': 
            pdf_list = self.split_pdf(self.path) # проверяем кол-во страниц в pdf
            if pdf_list: # если больше одной страницы возвращается список
                text = ''
                for pdf_path in pdf_list:
                    if os.path.exists(pdf_path):
                        text += await self.extract_text_from_pdf(pdf_path, 'pdf') # получаем текст из pdf
                        os.remove(pdf_path) # Удаление временного файла после его использования
            else: # если одна страница
                text = await self.extract_text_from_pdf(self.path, self.extension) # получаем текст из pdf
        elif self.extension == 'jpg':
            text = await self.extract_text_from_pdf(self.path, self.extension) # получаем текст из изображения
        else: return None
        if not text: return None
        # print('Получен текст:')
        # print(text)
        # print('---------------------------------------------------------')
        data = await self.get_info_docs(text) # получаем ключевые данные из текста
        info_docs = data.replace('**', '')
        info = info_docs.split('\n')
        for i in info: # заполняем ключевые данные из текста
            for key in data_company.keys():
                if key in i or (key.lower() in i and key == 'Телефон'):
                    if len(i.split(':')) > 1: 
                        value = i.split(':')[1].strip().replace(',', '').replace('"', '').replace('[', '').replace(']', '')
                        if key == 'Телефон': value = self.find_number(value)
                        if key == 'Наименование организации': value = value.replace('\\', '') 
                        if key == 'Расчетный счет' or key == 'Корреспондентский счет':
                            value = self.find_account(value)
                        if data_company[key] == '': 
                            data_company[key] = value
                        else:
                            data_company[key] += '|' + value
        # print('Информация о компании:')
        # print(data_company)
        # print('---------------------------------------------------------')
        count_companies = max(len(value.split('|')) for value in data_company.values()) # кол-во компаний данные о которых получены
        company_dicts = [{} for _ in range(count_companies)] if count_companies > 1 else [data_company]
        for key in data_company.keys():
            values = data_company[key].split('|')
            for i in range(len(company_dicts)): # Обрабатываем значения для каждой компании
                try:
                    company_dicts[i][key] = values[i]
                except IndexError:
                    company_dicts[i][key] = ''  # Если значений меньше чем компаний оставляем пустое значение
        final_dict = {}
        for i, company in enumerate(company_dicts):
            full_data = self.get_dadata(company) # получаем остальные данные о компании через Dadata
            if full_data: # если данные о компании найдены
                full_data['buyer']['user_id'] = self.chat_id # добавляем chat_id пользователя
                full_data[i] = full_data['buyer']
                del full_data["buyer"]
                final_dict.update(full_data)
        if final_dict:
            return final_dict
        else:
            return None

    # получаем ключевые данные из текста
    async def get_info_docs(self, text):
        response = await self.get_token(self.SECRET_TOKEN)
        if response == 1: return None
        access_token = response['access_token']
        url = "https://gigachat.devices.sberbank.ru/api/v1/chat/completions"
        message = f'''На основе предоставленного текста с данными о компании, составьте словари в формате json, используя следующие ключи:
        'Наименование организации': ''
        'ИНН': '' # 10 цифр
        'БИК': '' # 9 цифр
        'Расчетный счет': '' # Только номер российского расчетного счета (20 цифр)
        'Корреспондентский счет': '' # Только номер российского Корреспондентского счета (20 цифр)
        'Телефон': ''
        'Электронная почта': ''
        'Основание актов': '' # Устав; Свидетельство; Доверенность
        'Город': ''
        'Адрес': ''
        'Имя': ''
        'Фамилия': ''
        'Отчество': ''

        Если какая-то информация отсутствует в тексте, оставьте соответствующее значение ключа пустым.
        ВАЖНО: Если присутствует информация о нескольких компаниях укажите их все, каждую компанию в отельном словаре. 
        Если компания одна, в ответ отправить один словарь!
        
        Текст: {text}'''
        payload = json.dumps({
                "model": "GigaChat-Pro",
                "messages": [
                    {
                    "role": "user",
                    "content": message
                    }
    ],
                "stream": False,
                "update_interval": 0
                })
        headers = {
            'Content-Type': 'application/json',
            'Authorization': f"Bearer {access_token}"
        }
        try:
            async with aiohttp.ClientSession() as session:
                async with session.post(url, headers=headers, data=payload, ssl=False) as resp:
                    response_json = await resp.json()
                    return response_json['choices'][0]['message']['content']
        except Exception as e:
            print(f"Ошибка: {e}")
            
    # получаем остальные данные о компании через Dadata (по ИНН или названию компании)
    def get_dadata(self, data):
        full_data_company = {"buyer": { # формат записи
            "user_id": "",
            "city": "",
            "organization_fullname": "",
            "organization_shortname": "",
            "inn": "",
            "ogrn": "",
            "kpp": "",
            "address": "",
            "phone_number": "",
            "email": "",
            "bic": "",
            "bank_info": "",
            "checking_account": "",
            "corporate_account": "",
            "first_name": "", 
            "second_name": "",
            "surname": "",
            "position_user": "",
            "acts_basis": "",
            "number_proxy": ""
        }}
        token = "84608b23de93259f56cfb64093818b2e0edc0bfd"
        dadata = Dadata(token)
        company_name = data['Наименование организации']
        inn = data['ИНН']
        bic = data['БИК']
        checking_account = data['Расчетный счет']
        corporate_account = data['Корреспондентский счет']
        phone_number = data['Телефон']
        email = data['Электронная почта']
        acts_basis = data['Основание актов']
        city = data['Город']
        address = data['Адрес']
        first_name = data['Имя']
        second_name = data['Фамилия']
        surname = data['Отчество']
        results = []
        if inn:
            results = dadata.find_by_id("party", inn) # получаем данные о компании по ИНН
            if not results:
                if company_name:
                    results = dadata.suggest("party", company_name) # получаем данные о компании по названию
        else:
            if company_name:
                results = dadata.suggest("party", company_name) # получаем данные о компании по названию
            
        if results:
            data_company = results[0]['data']
            extracted_data = self.extract_data(data_company) # извлекаем данные полученные от Dadata
            
            if extracted_data.get('full_name') and len(extracted_data['full_name'].split()) > 2: # записываем ФИО
                list_name = extracted_data['full_name'].split()
                full_data_company['buyer']['first_name'] = list_name[1]
                full_data_company['buyer']['second_name'] = list_name[0]
                full_data_company['buyer']['surname'] = list_name[2]

            for key in ['city', 'organization_fullname', 'organization_shortname', 'inn', 'ogrn', 'kpp', 'address', 'phone_number', 'email', 'position_user']:
                if extracted_data.get(key): # записываем данные полученные от Dadata
                    full_data_company['buyer'][key] = extracted_data[key]
                    
            fields_to_update = {
                'phone_number': phone_number,
                'email': email,
                'city': city,
                'address': address,
                'first_name': first_name,
                'second_name': second_name,
                'surname': surname
            }
            for field, value in fields_to_update.items(): # проверяем наличие данных и если нет, то записываем из документа
                if not full_data_company['buyer'].get(field) and value:
                    full_data_company['buyer'][field] = value
        else:
            full_data_company['buyer']['phone_number'] = phone_number if phone_number else '' # записываем данные полученные из документа
            full_data_company['buyer']['email'] = email if email else '' 
            full_data_company['buyer']['city'] = city if city else '' 
            full_data_company['buyer']['address'] = address if address else '' 
            full_data_company['buyer']['first_name'] = first_name if first_name else '' 
            full_data_company['buyer']['second_name'] = second_name if second_name else '' 
            full_data_company['buyer']['surname'] = surname if surname else '' 
            
        result = dadata.find_by_id("bank", bic) # Получаем данные о банке через БИК
        if result:
            full_data_company['buyer']['bank_info'] = result[0]['unrestricted_value']
            full_data_company['buyer']['bic'] = bic # Записываем 
        full_data_company['buyer']['checking_account'] = checking_account if checking_account else '' # записываем данные полученные из документа
        full_data_company['buyer']['corporate_account'] = corporate_account if corporate_account else ''
        full_data_company['buyer']['acts_basis'] = acts_basis if acts_basis else ''
        return full_data_company

# ----------------------------------------Получение токенов гигачат и яндекса---------------------------
    async def get_token(self, auth_token, scope='GIGACHAT_API_B2B'):
        # Создаем идентификатор UUID (36 знаков)
        rq_uid = str(uuid.uuid4())

        # API URL
        url = "https://ngw.devices.sberbank.ru:9443/api/v2/oauth"

        # Заголовки
        headers = {
            'Content-Type': 'application/x-www-form-urlencoded',
            'Accept': 'application/json',
            'RqUID': rq_uid,
            'Authorization': f'Basic {auth_token}'
        }

        # Тело запроса
        payload = {
            'scope': scope
        }

        async with aiohttp.ClientSession() as session:
            try:
                # Делаем POST запрос с отключенной SSL верификацией
                # (можно скачать сертификаты Минцифры, тогда отключать проверку не надо)
                async with session.post(url, headers=headers, data=payload, ssl=False) as response:
                    return await response.json()
            except Exception as e:
                print(f"Ошибка: {str(e)}")
                return 1 

    # получаем токен IAM яндекса
    async def get_iamToken(self, oauth_token):
        data = {
            "yandexPassportOauthToken": oauth_token
        }
        async with aiohttp.ClientSession() as session:
            async with session.post(
                "https://iam.api.cloud.yandex.net/iam/v1/tokens",
                json=data) as response:
                
                if response.status == 200:
                    answer = await response.json()
                    return answer['iamToken']
                else:
                    print("Ошибка при отправке запроса. Код ответа:", response.status)
                    print("Текст ответа:", await response.text())
                    return None

# ----------------------------------------UTILS---------------------------------------------------
    # кодируем файл и возвращаем результат.
    def encode_file(self, file_path):
        with open(file_path, "rb") as fid:
            file_content = fid.read()
        return base64.b64encode(file_content).decode("utf-8")

    # Извлекаем текст из DOCX файла
    def extract_text_from_docx(self, docx_path):
        doc = docx.Document(docx_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables: 
            rows = table.rows
            for row in rows:
                row_text = ''
                for cell in row.cells:
                    row_text += f'{cell.text.strip()} '
                text += f'{row_text}\n'
        return text.strip()

    # Извлекаем текст из PDF файла или изображения
    async def extract_text_from_pdf(self, path, extension):
        encode = self.encode_file(path)
        mimeType = "application/pdf" if extension == 'pdf' else "JPEG"
        data = {
            "mimeType": mimeType,
            "languageCodes": ["*"],
            "content": encode
        }
        url = "https://ocr.api.cloud.yandex.net/ocr/v1/recognizeText"
        iamToken = await self.get_iamToken(self.oauth_token)
        if iamToken == None:
            print('iamToken')
            return
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {iamToken}",
            "x-folder-id": self.folder_id,
            "x-data-logging-enabled": "true"
        }
        async with aiohttp.ClientSession() as session:
            try:
                async with session.post(url, headers=headers, json=data) as response:
                    result = await response.json()
                    info = result['result']["textAnnotation"]["fullText"]
                    return info
            except Exception as e:
                print(f"Ошибка при отправке запроса: {e}")
                return None

    # разбираем полученные данные о компании
    def extract_data(self, data_company):
        management_data = data_company.get('management') or {}
        return {
            'city': data_company.get('address', {}).get('data', {}).get('city', ''),
            'organization_fullname': data_company.get('name', {}).get('full_with_opf', ''),
            'organization_shortname': data_company.get('name', {}).get('short_with_opf', ''),
            'inn': data_company.get('inn', ''),
            'ogrn': data_company.get('ogrn', ''),
            'kpp': data_company.get('kpp', ''),
            'address': data_company.get('address', {}).get('unrestricted_value', ''),
            'phone_number': data_company.get('phones', []),
            'email': data_company.get('emails', []),
            'full_name': management_data.get('name', ''),
            'position_user': management_data.get('post', '')
        } 

    def find_number(self, text):
        text_search = text.replace('-', '').replace(' ', '').replace('(', '').replace(')', '')
        pattern = r'\b([9]\d{9}(?!\d)|[7-8]\d{10}(?!\d))'
        matches = re.findall(pattern, text_search)
        if matches:
            return matches[0]
        else:
            matches = re.findall(pattern, text)
            if matches:
                return matches[0]
            else:
                return ''

    def find_account(self, text):
        result = ''.join([char for char in text if char.isdigit()])
        if result and len(result) == 20:
            return result
        else: 
            return ''

    def split_pdf(self, input_file):
        reader = PdfReader(input_file) # Открываем исходный PDF-файл для чтения
        num_pages = len(reader.pages) # Получаем количество страниц
        output_filenames = []
        if num_pages > 1:
            for page_num in range(num_pages):
                writer = PdfWriter()
                writer.add_page(reader.pages[page_num]) # Добавляем текущую страницу в новый PDF-файл
                
                output_filename = f"{input_file[:-4]}_page_{page_num + 1}.pdf" # Сохраняем каждый файл отдельно
                output_filenames.append(output_filename)
                with open(output_filename, "wb") as out_file:
                    writer.write(out_file)
            return output_filenames
        else:
            return output_filenames